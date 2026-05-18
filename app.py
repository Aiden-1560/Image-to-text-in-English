import streamlit as st
from docx import Document
from docx.shared import Pt
from io import BytesIO
import time
import re
from google import genai
from google.genai import types
from google.genai.errors import APIError

# 1. 페이지 기본 설정 및 디자인
st.set_page_config(
    page_title="Image To Text in English",
    page_icon="📝",
    layout="centered"
)

# 세련되고 직관적인 커스텀 CSS
st.markdown("""
    <style>
    .stApp { background-color: #FDFBF6; }
    .main-title {
        font-size: 42px !important; 
        font-weight: 800;
        color: #5C715E;
        text-align: center;
        margin-bottom: 10px;
        letter-spacing: -1.5px;
        white-space: nowrap;
    }
    .sub-title {
        font-size: 18px;
        color: #7A7A7A;
        text-align: center;
        margin-bottom: 45px;
        font-weight: 400;
        letter-spacing: -0.5px;
    }
    .stFileUploader {
        border-radius: 15px !important;
        background-color: white;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        padding: 10px;
    }
    div.stButton > button:first-child {
        background-color: #94A69A;
        color: white;
        border-radius: 10px;
        border: none;
        padding: 12px 40px;
        font-weight: 600;
        font-size: 16px;
        width: 100%;
        transition: all 0.3s ease;
    }
    div.stButton > button:first-child:hover { background-color: #7E8F83; }
    .status-box {
        padding: 25px;
        border-radius: 20px;
        background-color: white;
        margin-top: 30px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        text-align: center;
    }
    .percent-text {
        font-size: 20px;
        font-weight: 700;
        color: #5C715E;
        margin-bottom: 8px;
        text-align: left;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown('<p class="main-title">Image To Text in English</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-title">사진 속 지문을 인식하여 편집 가능한 워드 문서(.docx)로 변환합니다.</p>', unsafe_allow_html=True)

# 2. API Key 자동 로드
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    client = genai.Client(api_key=api_key)
    
    # 3. 파일 업로드 UI
    uploaded_files = st.file_uploader(
        "변환할 영어 지문 사진을 업로드하세요 (복수 선택 가능)", 
        type=["jpg", "jpeg", "png"],
        accept_multiple_files=True
    )

    if uploaded_files:
        st.write(f"📂 **{len(uploaded_files)}개**의 파일이 선택되었습니다.")
        
        if st.button("Word 파일로 변환하기 ✨"):
            
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
            percent_display = st.empty()  
            progress_bar = st.progress(0)  
            status_text = st.empty()       
            
            total_files = len(uploaded_files)
            model_name = 'gemini-2.5-flash'
            
            # 💡 초정밀 흐름 제어를 위한 시간 기반 변수 설정
            # 사진 1장당 처리 예상 시간 가이드라인 설정 (API연동 + 대기시간 포함 약 12초)
            estimated_seconds_per_file = 12 
            total_estimated_seconds = total_files * estimated_seconds_per_file
            
            # 실시간 가상 퍼센트를 올릴 스레드 타이머 실행
            current_percent = 0
            
            for idx, file in enumerate(uploaded_files):
                image_bytes = file.read()
                
                # 원장님 피드백 완벽 반영 프롬프트 (대화 주체 분리 및 소제목 구조화 요청)
                prompt = """
                이 사진 속의 영어 지문 텍스트를 정확하게 추출해줘.
                - 사진의 메인 제목이나 소제목이 있다면 텍스트 제일 앞에 '[HEADING]' 이라는 태그를 붙여줘. (예: [HEADING] Reports from the Battlefield)
                - 대화 내용(Dialogue) 구조인 경우, 대화의 주체를 나타내는 이름 뒤에 콜론(:)을 붙이고, 이름 앞에 '[NAME]' 태그를 붙여줘. (예: [NAME] Mike: Hey, guys!)
                - 일반 본문 문장은 아무 태그 없이 문맥에 맞게 가독성 높은 단락으로 나누어줘.
                - 원본에서 **진하게** 처리된 강조 키워드가 있다면 마크다운 기호(**)를 유지해줘.
                - 결과물은 오직 추출된 텍스트만 보여주고, 다른 부연 설명은 하지 마.
                """
                
                # 파일 처리 시작하면서 퍼센트를 목표 지점까지 부드럽게 롤링 업
                target_start_percent = int((idx / total_files) * 100)
                target_end_percent = int(((idx + 0.5) / total_files) * 100)
                
                status_text.text(f"⏳ [{idx+1}/{total_files}] '{file.name}' 지문을 AI 서버에서 분석하는 중...")
                
                # API 호출 전에 게이지를 자연스럽게 먼저 밀어줌
                for p in range(current_percent, target_end_percent):
                    current_percent = p
                    percent_display.markdown(f'<p class="percent-text">⏳ 변환 진행률: {current_percent}%</p>', unsafe_allow_html=True)
                    progress_bar.progress(current_percent)
                    time.sleep(0.04)
                
                extracted_text = ""
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=[
                            types.Part.from_bytes(data=image_bytes, mime_type=file.type),
                            prompt
                        ]
                    )
                    extracted_text = response.text
                    
                except APIError as e:
                    if e.code == 429:
                        status_text.text("⏳ 처리하는데 시간이 걸리니 조금만 기다려주세요... (잠시 대기 중)")
                        time.sleep(12)
                        response = client.models.generate_content(
                            model=model_name,
                            contents=[types.Part.from_bytes(data=image_bytes, mime_type=file.type), prompt]
                        )
                        extracted_text = response.text
                    else:
                        st.error(f"❌ '{file.name}' 처리 중 오류 발생: {str(e)}")
                        continue

                # 텍스트 추출 성공 시 워드 문서 조립 및 6초 대기 구간 게이지 연동
                if extracted_text:
                    # 파일 완료 시점 목표 퍼센트 계산
                    file_done_percent = int(((idx + 1) / total_files) * 100)
                    if idx == total_files - 1:
                        file_done_percent = 100
                    
                    # 워드 파일 타이틀 추가
                    doc.add_heading(f"Source: {file.name}", level=3)
                    
                    paragraphs = extracted_text.split('\n')
                    for para_text in paragraphs:
                        clean_text = para_text.strip()
                        if not clean_text:
                            continue
                        
                        p_tag = doc.add_paragraph()
                        
                        # 케이스 1: 제목/소제목인 경우 -> 전체 크고 진하게
                        if clean_text.startswith("[HEADING]"):
                            heading_content = clean_text.replace("[HEADING]", "").strip()
                            run = p_tag.add_run(heading_content)
                            run.bold = True
                            run.font.size = Pt(14)
                            
                        # 케이스 2: 대화문 구조인 경우 -> 주체(이름)만 골라 진하게
                        elif clean_text.startswith("[NAME]"):
                            name_content = clean_text.replace("[NAME]", "").strip()
                            # '이름:' 과 '나머지 대화내용' 분리 추출
                            match = re.match(r"^([^:]+:)(.*)$", name_content)
                            if match:
                                name_part = match.group(1)   # 이름:
                                dialogue_part = match.group(2) # 대화 내용
                                
                                # 이름 부분 볼드 처리
                                r_name = p_tag.add_run(name_part)
                                r_name.bold = True
                                
                                # 대화 내용 부분은 일반 서체 처리
                                p_tag.add_run(dialogue_part)
                            else:
                                p_tag.add_run(name_content)
                                
                        # 케이스 3: 일반 지문 본문 문장
                        else:
                            # 일반 본문 내에 부분 강조(**)가 섞여 있을 경우 처리
                            parts = clean_text.split('**')
                            for i, part in enumerate(parts):
                                run = p_tag.add_run(part)
                                if i % 2 == 1:
                                    run.bold = True
                                    
                    doc.add_page_break()
                    
                    # 💡 [점프 현상 차단 패치] 다음 파일로 넘어가기 전 안전 휴식(6초) 시간 동안
                    # 퍼센트 바가 멈추지 않고 1%씩 째깍째깍 올라가도록 연동시킵니다.
                    if idx < total_files - 1:
                        steps = 60 # 6초를 60프레임으로 쪼개서 실행
                        percent_increment = (file_done_percent - current_percent) / steps
                        
                        for step in range(steps):
                            current_percent += percent_increment
                            display_p = min(int(current_percent), file_done_percent)
                            
                            percent_display.markdown(f'<p class="percent-text">⏳ 변환 진행률: {display_p}%</p>', unsafe_allow_html=True)
                            progress_bar.progress(display_p)
                            
                            # 카운트다운 초 매핑 연출 (10프레임마다 1초씩 차감)
                            sec_left = 6 - (step // 10)
                            status_text.text(f"⏳ 처리하는데 시간이 걸리니 조금만 기다려주세요.. ({sec_left}초)")
                            time.sleep(0.1)
                            
                    current_percent = file_done_percent
                    percent_display.markdown(f'<p class="percent-text">⏳ 변환 진행률: {current_percent}%</p>', unsafe_allow_html=True)
                    progress_bar.progress(current_percent)
            
            # 최종 100% 도달 완료
            percent_display.markdown('<p class="percent-text" style="color:#0D9488;">🎉 변환 진행률: 100%</p>', unsafe_allow_html=True)
            progress_bar.progress(100)
            status_text.text("🎉 모든 영어 지문이 성공적으로 변환되었습니다!")
            
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            st.markdown('<div class="status-box">', unsafe_allow_html=True)
            st.download_button(
                label="📥 변환된 Word 파일 다운로드",
                data=docx_buffer,
                file_name="Converted_English_Texts.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.markdown('</div>', unsafe_allow_html=True)

except KeyError:
    st.error("🔒 설정 오류: Streamlit Secrets에 API Key를 등록해 주세요.")
